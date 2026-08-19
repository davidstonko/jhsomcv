# -*- coding: utf-8 -*-
"""Generate assets/jhsom-cv-template.docx: the full ABMF skeleton in the departmental Word geometry."""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from taxonomy import TAXONOMY
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT

FONT='Times New Roman'; SZ=Pt(11)
doc=Document()
st=doc.styles['Normal']; st.font.name=FONT; st.font.size=SZ
pf=st.paragraph_format; pf.space_before=Pt(0); pf.space_after=Pt(0); pf.line_spacing=1.0
s=doc.sections[0]
s.page_width=Inches(8.5); s.page_height=Inches(11)
s.left_margin=Inches(0.75); s.right_margin=Inches(0.5)
s.top_margin=Inches(0.75); s.bottom_margin=Inches(0.5)

def P(text='', bold=False, align=None, left=None, first=None, tabs=()):
    p=doc.add_paragraph()
    r=p.add_run(text); r.bold=bold; r.font.name=FONT; r.font.size=SZ
    f=p.paragraph_format
    if align is not None: p.alignment=align
    if left is not None: f.left_indent=Inches(left)
    if first is not None: f.first_line_indent=Inches(first)
    for t in tabs: f.tab_stops.add_tab_stop(Inches(t))
    return p

from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
P('CURRICULUM VITAE', bold=True, align=AL.CENTER)
P('The Johns Hopkins University School of Medicine', align=AL.CENTER)
P(); P(); P()
P('<Month D, YYYY>')
P('<First M. Last, degrees>\t\t\t\t\t\tDate of this version')
P()
levels=[l for l,_,_ in TAXONOMY]
for i,(lvl, title, note) in enumerate(TAXONOMY):
    leaf = (i+1>=len(TAXONOMY)) or (levels[i+1] <= lvl)
    if lvl==1: P(); P(title)
    elif lvl==2: P(); P(title)
    elif lvl==3: P(title)
    else: P('\t'+title)
    if leaf:
        P('NONE' if lvl>1 else 'NONE')
doc.save(os.path.join(os.path.dirname(__file__),'..','assets','jhsom-cv-template.docx'))
print('written')
